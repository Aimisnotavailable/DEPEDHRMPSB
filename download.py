# download.py
import io, json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def download_pdf(
    code,
    participant,
    evaluation,
    interview,
    validation,
    EVAL_STRUCTURE,
    header_image_path: str = "./static/images/header_docx.png",
    logo_paths: list = None
):
    # 1) Load applicant + extra_data
    applicant = participant.query.get_or_404(code)
    try:
        extra_data = json.loads(applicant.extra_data or "{}")
    except:
        extra_data = {}

    # 2) Fetch evaluation records
    eval_records = evaluation.query.filter_by(
        interview_id=applicant.interview_id,
        participant_code=applicant.code
    ).all()

    # 3) Determine eval type & compute NCOI if teaching
    eval_type   = interview.query.get(applicant.interview_id).type
    eval_struct = EVAL_STRUCTURE[eval_type]
    scores, ncoi = [], None

    if eval_type == "teaching" and eval_records:
        for rec in eval_records:
            ed = json.loads(rec.extra_data or "{}")
            total = 0
            for crit, fields in eval_struct.items():
                for f in fields:
                    total += ed.get(crit, {}).get(f, 0)
            scores.append(round(total, 2))
        avg = sum(scores) / len(scores) if scores else 0
        try:
            trf = float(extra_data.get("trf_rating", 0))
        except:
            trf = 0
        ncoi = round(trf + avg, 2)

    # default logo paths if none provided
    if logo_paths is None:
        logo_paths = [
            "./static/images/deped_seal.png",
            "./static/images/deped_logo.png",
            "./static/images/tagalog_deped.png",
        ]

    # 4) Build the document
    doc     = Document()
    section = doc.sections[0]

    # compute usable width inside margins
    avail_w = (
        section.page_width
        - section.left_margin
        - section.right_margin
    )

    # ---- HEADER: full-width banner image ----
    hdr = section.header.add_paragraph().add_run()
    hdr.width = avail_w
    hdr.add_picture(header_image_path, width=avail_w)

    # ---- BODY: Title + Applicant Details ----
    title = doc.add_heading(f'Applicant {applicant.code} Details', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text = 'Field', 'Information'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    data_rows = [
        ('Name', applicant.name or ''),
        ('Address', applicant.address or ''),
        ('Birthday', str(applicant.birthday) if applicant.birthday else ''),
        ('Age', str(applicant.age) if applicant.age else ''),
        ('Sex', applicant.sex or ''),
        ('Raw Education', str(applicant.raw_edu) or ''),
        ('Raw Experience', str(applicant.raw_exp) or ''),
        ('Raw Training', str(applicant.raw_trn) or ''),
        ('Score Education', str(applicant.score_edu) or ''),
        ('Score Experience', str(applicant.score_exp) or ''),
        ('Score Training', str(applicant.score_trn) or '')
    ]
    if eval_type == 'teaching':
        data_rows += [
            ('LPT/PBET/LEPT Raw Score', str(extra_data.get('lpt_rating', ''))),
            ('COI',                    str(extra_data.get('COI', ''))),
            ('TRF',                    str(extra_data.get('trf_rating', '')))
        ]
        if ncoi is not None:
            data_rows.append(('NCOI', str(ncoi)))

    for fld, val in data_rows:
        row = tbl.add_row().cells
        row[0].text, row[1].text = fld, val
    
     # ---- ATTESTATION ----
    doc.add_paragraph()
    att = doc.add_paragraph(
        "I hereby attest to the conduct of the application and assessment process "
        "in accordance with the applicable guidelines."
    )
    att.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- THREE LOGOS INLINE (no table) ----
    doc.add_paragraph()  # spacer
    logos_para = doc.add_paragraph()
    logos_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_w = Inches(1.1)  # smaller width for each logo

    for path in logo_paths:
        run = logos_para.add_run()
        run.add_picture(path, width=logo_w)
        logos_para.add_run('  ')  # space between logos

    # ---- Evaluations / Criteria ----
    doc.add_paragraph()
    if eval_type == 'teaching' and scores:
        doc.add_heading('Evaluations', level=2)
        for i, sc in enumerate(scores, 1):
            doc.add_paragraph(f'Evaluator #{i} score: {sc}', style='List Bullet')
    elif eval_records:
        doc.add_heading('Evaluation Criteria', level=2)
        et = doc.add_table(rows=1, cols=4)
        et.style = 'Table Grid'
        hdr = et.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
            'Criteria', 'Weight', 'Qualification', 'Score'
        )
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        qual_map = {
            'Education':  applicant.raw_edu,
            'Experience': applicant.raw_exp,
            'Training':   applicant.raw_trn
        }
        ev0 = json.loads(eval_records[0].extra_data or "{}")
        for crit, cfg in eval_struct.items():
            row = et.add_row().cells
            row[0].text = crit
            row[1].text = str(cfg.get('weight', ''))
            row[2].text = str(
                qual_map.get(crit, ev0.get(crit, {}).get('qualification', 'N/A'))
            )
            row[3].text = str(round(ev0.get(crit, 0), 2))


    # ---- FOOTER ----
    footer = section.footer
    fp     = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr     = fp.add_run(
        "laoag.city@deped.gov.ph | facebook.com/depedtayolaoagcity | depedlaoagcity.com"
    )
    fr.font.size = Pt(9)

    # 5) Save to memory and return
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
