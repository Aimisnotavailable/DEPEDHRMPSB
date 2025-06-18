import json, io
from flask import Flask, render_template, request, send_file
from openpyxl import load_workbook
from docx import Document

app = Flask(__name__)
HR_XLSX     = "hr_data.xlsx"
EVAL_JSON   = "evaluations.json"

# load/save evaluations from a JSON file
def load_evals():
    try:
        with open(EVAL_JSON, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}

def save_evals(e):
    with open(EVAL_JSON, "w") as f:
        json.dump(e, f, indent=2)

# read applicants basic list
def load_applicants():
    wb = load_workbook(HR_XLSX, data_only=True)
    ws = wb["Applicants"]
    hdr = {c.value:i for i,c in enumerate(ws[1])}
    return [
      {"id": row[hdr["ApplicantID"]], "name": row[hdr["Name"]]}
      for row in ws.iter_rows(min_row=2, values_only=True)
    ]

# read full row of applicant
def fetch_applicant(app_id):
    wb = load_workbook(HR_XLSX, data_only=True)
    ws = wb["Applicants"]
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if str(row[headers.index("ApplicantID")]) == str(app_id):
            return {hdr: row[i] for i, hdr in enumerate(headers)}
    return None

# compute qualification‐based score total (not shown on screen)
def compute_qual_scores(details):
    d = details
    scores = {
      "Education":   d["QS_Increment"],
      "Training":    d["TrainingHours"]//8,
      "Experience":  min((d["ExperienceMonths"]//12)*2.5,20),
      "Performance": d["PerformanceRating"]*(20/5),
      "Outstanding": d["OutstandingPoints"],
      "AppOfEdu":    d["AppOfEducationPts"],
      "AppOfLD":     d["AppOfLD_Pts"],
      "Potential":   d["PotentialPts"],
    }
    return scores, sum(scores.values())

@app.route("/", methods=["GET"])
def form():
    apps = load_applicants()
    sel  = request.args.get("applicant_id")
    details = fetch_applicant(sel) if sel else None

    evals = load_evals()
    user_scores = evals.get(sel, {
      "aptitude":0,"character_traits":0,"fitness":0,
      "leadership":0,"communication":0,"trf_rating":0
    })
    total_eval = sum(float(v) for v in user_scores.values())

    return render_template(
      "form.html",
      applicants=apps,
      details=details,
      user_scores=user_scores,
      total_eval=total_eval
    )

@app.route("/generate", methods=["POST"])
def generate():
    app_id = request.form["applicant_id"]
    details = fetch_applicant(app_id)
    if not details:
        return "Applicant not found", 404

    # grab & save the six evaluation fields
    fields = ["aptitude","character_traits","fitness",
              "leadership","communication","trf_rating"]
    evals = load_evals()
    evals[app_id] = {f: float(request.form[f]) for f in fields}
    save_evals(evals)

    # recalc totals
    _, qual_total = compute_qual_scores(details)
    user = evals[app_id]
    eval_total = sum(user.values())

    # build doc
    doc = Document()
    doc.add_heading("INDIVIDUAL EVALUATION SHEET (IES)", level=1)

    # Basic info
    basic = [("Name of Applicant", details["Name"]),
             ("Application Code",   details["AppCode"]),
             ("Position Applied For", details["Position"]),
             ("Schools Division Office", details["SDO"]),
             ("Contact Number",     details["Contact"]),
             ("Job Group/SG-Level", details["JobGroup"])]
    tbl1 = doc.add_table(rows=len(basic), cols=2)
    tbl1.style = "Light List"
    for i,(lbl,val) in enumerate(basic):
        r = tbl1.rows[i]
        r.cells[0].text = lbl
        r.cells[1].text = str(val)

    doc.add_paragraph()
    # Qualifications
    qual_keys = ["QS_Increment","TrainingHours","ExperienceMonths",
                "PerformanceRating","OutstandingPoints",
                "AppOfEducationPts","AppOfLD_Pts","PotentialPts"]
    tbl2 = doc.add_table(rows=1+len(qual_keys), cols=2)
    tbl2.style = "Light List Accent 1"
    tbl2.rows[0].cells[0].text = "Criterion"
    tbl2.rows[0].cells[1].text = "Value"
    for i,key in enumerate(qual_keys,1):
        r = tbl2.rows[i]
        r.cells[0].text = key.replace("_"," ")
        r.cells[1].text = str(details[key])

    doc.add_paragraph()
    # Evaluation breakdown
    tbl3 = doc.add_table(rows=1+len(fields)+1, cols=3)
    tbl3.style = "Light List Accent 1"
    hdr = tbl3.rows[0].cells
    hdr[0].text = "Criterion"
    hdr[1].text = "Score"
    hdr[2].text = "Max Score"

    maxes = {"aptitude":1,"character_traits":1,"fitness":1,
             "leadership":1,"communication":1,"trf_rating":20}
    for f in fields:
        row = tbl3.add_row().cells
        row[0].text = f.replace("_"," ").title()
        row[1].text = f"{evals[app_id][f]:.2f}"
        row[2].text = f"{maxes[f]:.2f}"
    total_row = tbl3.add_row().cells
    total_row[0].text = "TOTAL QUAL"
    total_row[1].text = f"{qual_total:.2f}"
    total_row[2].text = ""
    final_row = tbl3.add_row().cells
    final_row[0].text = "TOTAL EVAL"
    final_row[1].text = f"{eval_total:.2f}"
    final_row[2].text = "25.00"

    doc.add_paragraph("\nName & Signature of Applicant: ____________________   Date: ________")
    doc.add_paragraph("Attested:\nMARIECON G. RAMIREZ EdD, CESO VI\nHRMPSB Chair")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return send_file(bio,
                     as_attachment=True,
                     download_name=f"IES_{details['AppCode']}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(debug=True)
