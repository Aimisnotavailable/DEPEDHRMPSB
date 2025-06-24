from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from flask import app, send_file
import json

from app import EVAL_STRUCUTURE, Evaluation, Interview, Participant, Validation, admin_required

# Add this new route alongside your existing applicant_detail route
@app.route("/admin/applicant/<code>/download-pdf")
@admin_required
def download_applicant_pdf(code):
    applicant = Participant.query.get_or_404(code)
    extra_data = None
    if applicant.extra_data:
        try:
            extra_data = json.loads(applicant.extra_data)
        except Exception:
            extra_data = applicant.extra_data

    validations = Validation.query.filter_by(
        interview_id=applicant.interview_id,
        participant_code=applicant.code
    ).all()

    eval_records = Evaluation.query.filter_by(
         interview_id=applicant.interview_id,
         participant_code=applicant.code
    ).all()

    # For teaching interviews, compute NCOI as:
    #   NCOI = Admin's TRF (max 20) + Average of evaluators' five criteria scores (max 5)
    eval_type = Interview.query.filter_by(id=applicant.interview_id).first().type
    eval_struct = EVAL_STRUCUTURE[eval_type]

    ncoi = None
    avg_eval = None
    scores = []
    if applicant.interview.type == "teaching" and eval_records:
        
        for eval_record in eval_records:
            overall = 0
            extra_data_json = json.loads(eval_record.extra_data)
            for key in eval_struct.keys():
                for field in eval_struct[key].keys():
                    overall = round(overall + extra_data_json[key][field], 2)
                scores.append(overall)

        total_eval = sum(score for score in scores)
        avg_eval = total_eval / len(eval_records)
        try:
            admin_trf = float(json.loads(applicant.extra_data).get("trf_rating", 0))
        except Exception:
            admin_trf = 0
        ncoi = round(admin_trf + avg_eval, 2)

    # Create Word document
    doc = Document()
    
    # Add header with logo and title
    header_paragraph = doc.add_paragraph()
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_paragraph.add_run("SCHOOLS DIVISION OF LAOAG CITY")
    run.bold = True
    run.font.size = Inches(0.2)
    
    doc.add_paragraph()  # Add space
    
    # Add title
    title = doc.add_heading(f'Applicant {applicant.code} Details', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create main details table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Add table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    # Make headers bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add applicant data rows
    data_rows = [
        ('Name', applicant.name or ''),
        ('Address', applicant.address or ''),
        ('Birthday', str(applicant.birthday) if applicant.birthday else ''),
        ('Age', str(applicant.age) if applicant.age else ''),
        ('Sex', applicant.sex or ''),
        ('Raw Education', str(applicant.raw_edu) if applicant.raw_edu else ''),
        ('Raw Experience', str(applicant.raw_exp) if applicant.raw_exp else ''),
        ('Raw Training', str(applicant.raw_trn) if applicant.raw_trn else ''),
        ('Score Education', str(applicant.score_edu) if applicant.score_edu else ''),
        ('Score Experience', str(applicant.score_exp) if applicant.score_exp else ''),
        ('Score Training', str(applicant.score_trn) if applicant.score_trn else ''),
    ]
    
    # Add teaching-specific data if applicable
    if applicant.interview.type == 'teaching' and extra_data:
        data_rows.extend([
            ('LPT/PBET/LEPT Raw Score', str(extra_data.get('lpt_rating', '')) if extra_data.get('lpt_rating') else ''),
            ('COI', str(extra_data.get('COI', '')) if extra_data.get('COI') else ''),
            ('TRF', str(extra_data.get('trf_rating', '')) if extra_data.get('trf_rating') else ''),
        ])
        if ncoi is not None:
            data_rows.append(('NCOI', str(ncoi)))
    
    # Add rows to table
    for field, value in data_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = value
    
    # Add evaluations section if scores exist
    if scores:
        doc.add_paragraph()  # Add space
        doc.add_heading('Evaluations', 2)
        
        eval_table = doc.add_table(rows=1, cols=2)
        eval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        eval_table.style = 'Table Grid'
        
        eval_hdr_cells = eval_table.rows[0].cells
        eval_hdr_cells[0].text = 'Evaluator'
        eval_hdr_cells[1].text = 'Score'
        
        # Make headers bold
        for cell in eval_hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add evaluation scores
        for i, score in enumerate(scores, 1):
            eval_row_cells = eval_table.add_row().cells
            eval_row_cells[0].text = f'Evaluator #{i}'
            eval_row_cells[1].text = str(round(score, 2))
    
    # Add footer
    doc.add_paragraph()
    footer_paragraph = doc.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_paragraph.add_run("laoag.city@deped.gov.ph | facebook.com/depedtayolaoagcity | depedlaoagcity.com")
    footer_run.font.size = Inches(0.12)
    
    # Save document to memory
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Return file for download
    return send_file(
        doc_io,
        as_attachment=True,
        download_name=f'applicant_{applicant.code}_details.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )