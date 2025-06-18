import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 1) LOAD EXCEL
xlsx_path = "sample_evaluation.xlsx"   # or your real file
xls = pd.ExcelFile(xlsx_path)

# 1a) Basic Info sheet → 2-column Label|Value
info_df = pd.read_excel(xls, sheet_name="Basic Info", header=None, index_col=0)
app_info = info_df[1].to_dict()

# 1b) Rating Sheet → must have these exact columns
scores_df = pd.read_excel(xls, sheet_name="Rating Sheet")

# sanity check
expected = {"Criteria", "Weight", "Details", "Computation"}
missing = expected - set(scores_df.columns)
if missing:
    raise KeyError(f"Your Rating Sheet is missing columns: {missing}")

# 2) PARSE AND ENSURE NUMERIC
scores_df["Weight"]      = pd.to_numeric(scores_df["Weight"],      errors="raise")
scores_df["Computation"] = pd.to_numeric(scores_df["Computation"], errors="raise")

total_wt  = scores_df["Weight"].sum()
total_pts = scores_df["Computation"].sum()

# 3) BUILD DOCX
doc = Document()

# 3a) DepEd header
hdr = doc.add_heading("", level=0)
hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
h = hdr.add_run(
    "Republic of the Philippines\n"
    "Department of Education\n"
    "Region I – Schools Division of Laoag City\n\n"
    "INDIVIDUAL EVALUATION SHEET (IES)\n\n"
)
h.bold = True

# 3b) Applicant basic info
fields = [
    "Name of Applicant", "Application Code",
    "Position Applied For", "Schools Division Office",
    "Contact Number", "Job Group/SG-Level"
]
tbl_info = doc.add_table(rows=0, cols=2)
tbl_info.style = "Light Shading Accent 1"
for f in fields:
    c0, c1 = tbl_info.add_row().cells
    c0.text = f
    c0.paragraphs[0].runs[0].bold = True
    c1.text = str(app_info.get(f, ""))

doc.add_paragraph()  # spacer

# 3c) Main evaluation table
cols = ["CRITERIA", "WEIGHT", "Details", "Computation"]
tbl = doc.add_table(rows=1, cols=len(cols), style="Table Grid")
for idx, title in enumerate(cols):
    cell = tbl.rows[0].cells[idx]
    cell.text = title
    cell.paragraphs[0].runs[0].bold = True

for _, row in scores_df.iterrows():
    c0, c1, c2, c3 = tbl.add_row().cells
    c0.text = str(row["Criteria"])
    c1.text = str(int(row["Weight"]))
    c2.text = str(row["Details"])
    c3.text = f"{row['Computation']:.2f}"

# Totals row
tot = tbl.add_row().cells
tot[0].text = "TOTAL"
tot[1].text = str(int(total_wt))
tot[2].text = ""
tot[3].text = f"{total_pts:.2f}"

doc.add_paragraph("\n")  # spacer

# 3d) Attestation block
p = doc.add_paragraph()
p.add_run(
    "I hereby attest to the conduct of the application and assessment process in accordance "
    "with applicable guidelines; and acknowledge the results of the comparative assessment "
    "and the points given based on my qualifications and submitted documentary requirements.\n\n"
).italic = True

sig = doc.add_paragraph()
sig.add_run("Name & Signature of Applicant: ").bold = True
sig.add_run("________________________    Date: __________\n")
sig.add_run("Attested by (HRMPSB Chair): ").bold = True
sig.add_run(f"{app_info.get('HRMPSB Chair','')}    Date: {datetime.today():%m/%d/%y}")

# 3e) Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run(
    "DepEd Laoag City · P. Gomez St., Brgy. 23 San Matias, Laoag City · "
    "Tel. (077) 771-3678 · laoag.city@deped.gov.ph"
).italic = True

# 4) SAVE
out_name = "IES_Automated.docx"
doc.save(out_name)
print(f"✅ {out_name} created.")
